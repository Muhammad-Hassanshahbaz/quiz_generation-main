from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
from docx import Document
from io import BytesIO
import os
import datetime

# Initialize FastAPI app
app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
client = Groq(api_key="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")

class QuizRequest(BaseModel):
    text: str
    num_questions: int = 10

def preprocess_text(text: str) -> str:
    """Pre-process text for question generation"""
    prompt = f"Normalize and simplify the following text for question generation:\n{text}"
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You are an expert in simplifying text for educational purposes."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

def select_sentences(text: str, num_questions: int) -> list:
    """Select important sentences for questions"""
    prompt = f"Select {num_questions} most important sentences from the following text for creating quiz questions:\n{text}"
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": f"You are an expert in identifying {num_questions} important sentences for quizzes."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=300
    )
    return response.choices[0].message.content.strip().split("\n")

def generate_question(sentence: str) -> str:
    """Generate MCQ from sentence"""
    prompt = f"Create a multiple-choice question from the following sentence:\n{sentence}"
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {"role": "system", "content": "You are an expert in generating quiz questions."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=200
    )
    return response.choices[0].message.content.strip()

def generate_quiz(text: str, num_questions: int) -> list:
    """Generate full quiz"""
    preprocessed_text = preprocess_text(text)
    sentences = select_sentences(preprocessed_text, num_questions)
    return [generate_question(sentence) for sentence in sentences]

def create_quiz_document(quiz: list) -> BytesIO:
    """Create DOCX in memory"""
    doc = Document()
    doc.add_heading('Generated Quiz', 0)
    
    for i, question in enumerate(quiz, 1):
        doc.add_paragraph(f"Question {i}: {question}")
        doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.post("/generate/")
async def generate_quiz_endpoint(request: QuizRequest):
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        if request.num_questions < 1:
            raise HTTPException(status_code=400, detail="Number of questions must be at least 1")
            
        quiz = generate_quiz(request.text, request.num_questions)
        return {"quiz": quiz}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download/")
async def download_quiz_endpoint(request: QuizRequest):
    try:
        quiz = generate_quiz(request.text, request.num_questions)
        buffer = create_quiz_document(quiz)
        
        filename = f"quiz_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/")
async def welcome():
    return {"message": "Welcome to Urdu Fashion Assistant API!"}


# For Vercel deployment
from mangum import Mangum
handler = Mangum(app)